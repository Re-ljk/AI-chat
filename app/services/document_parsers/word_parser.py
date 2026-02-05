"""
    @project: aihub
    @Author: jiangkuanli
    @file: word_parser
    @date: 2026/2/2
    @desc: Word文档解析器
"""

from typing import Dict, Any, List
import os
import base64
import tempfile
from pathlib import Path
from .base_parser import BaseDocumentParser


class WordParser(BaseDocumentParser):
    """Word文档解析器"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析Word文档
        
        Args:
            file_path: Word文档路径
            
        Returns:
            包含文档内容和元数据的字典
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            content = []
            images = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            full_text = '\n'.join(content)
            cleaned_text = self.clean_text(full_text)
            
            image_count = self._extract_images(doc, file_path, images)
            
            metadata = self.extract_metadata(cleaned_text)
            metadata.update({
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'image_count': image_count,
                'has_images': image_count > 0,
                'file_type': 'word'
            })
            
            return {
                'content': cleaned_text,
                'metadata': metadata,
                'images': images,
                'success': True
            }
            
        except ImportError:
            return {
                'content': '',
                'metadata': {},
                'images': [],
                'success': False,
                'error': 'python-docx库未安装，请运行: pip install python-docx'
            }
        except Exception as e:
            return {
                'content': '',
                'metadata': {},
                'images': [],
                'success': False,
                'error': f'解析Word文档失败: {str(e)}'
            }
    
    def _extract_images(self, doc, file_path: str, images: List[Dict[str, Any]]) -> int:
        """提取Word文档中的图片
        
        Args:
            doc: Document对象
            file_path: 文件路径
            images: 图片列表
            
        Returns:
            图片数量
        """
        try:
            image_count = 0
            
            for rel in doc.part.rels.values():
                if 'image' in rel.target_ref:
                    image_data = rel.target_part.blob
                    image_ext = rel.target_ref.split('.')[-1]
                    
                    image_info = {
                        'index': image_count,
                        'extension': image_ext,
                        'size': len(image_data),
                        'data': base64.b64encode(image_data).decode('utf-8')
                    }
                    
                    images.append(image_info)
                    image_count += 1
            
            return image_count
        except Exception as e:
            print(f"提取图片时出错: {str(e)}")
            return 0
    
    def parse_from_bytes(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        """从字节数据解析Word文档
        
        Args:
            file_content: 文件字节数据
            file_type: 文件类型
            
        Returns:
            包含文档内容和元数据的字典
        """
        try:
            from docx import Document
            
            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = Document(temp_file_path)
                content = []
                images = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text)
                
                full_text = '\n'.join(content)
                cleaned_text = self.clean_text(full_text)
                
                image_count = self._extract_images(doc, temp_file_path, images)
                
                metadata = self.extract_metadata(cleaned_text)
                metadata.update({
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'image_count': image_count,
                    'has_images': image_count > 0,
                    'file_type': 'word'
                })
                
                return {
                    'content': cleaned_text,
                    'metadata': metadata,
                    'images': images,
                    'success': True
                }
            finally:
                # 删除临时文件
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
            
        except ImportError:
            return {
                'content': '',
                'metadata': {},
                'images': [],
                'success': False,
                'error': 'python-docx库未安装，请运行: pip install python-docx'
            }
        except Exception as e:
            return {
                'content': '',
                'metadata': {},
                'images': [],
                'success': False,
                'error': f'解析Word文档失败: {str(e)}'
            }
    
    def get_supported_extensions(self) -> list:
        """获取支持的文件扩展名"""
        return ['.docx', '.doc']
